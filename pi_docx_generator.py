"""
pi_docx_generator.py — Proforma Invoice DOCX generator (Python/python-docx)
Replaces the previous Node.js approach.
"""

import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import lxml.etree as etree


# ── COLOUR CONSTANTS ──────────────────────────────────────────────────────────
C_DARK_GREEN  = "1B4332"
C_MID_GREEN   = "2D6A4F"
C_LIGHT_GREEN = "D8F3DC"
C_WHITE       = "FFFFFF"
C_DARK_TEXT   = "1A1A1A"
C_GRAY_TEXT   = "555555"
C_GOLD        = "B7950B"


# ── HELPER: SET CELL BACKGROUND ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def remove_cell_borders(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(tc_borders)


def add_run(para, text, bold=False, size_pt=9, color=None, italic=False):
    run = para.add_run(text or "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_paragraph_in_cell(cell, text, bold=False, size_pt=9, color=None,
                           align=WD_ALIGN_PARAGRAPH.LEFT, italic=False,
                           space_before=0, space_after=2):
    if len(cell.paragraphs) == 1 and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    add_run(p, text, bold=bold, size_pt=size_pt, color=color, italic=italic)
    return p


def section_header_row(table, text, col_span=None):
    """Add a full-width dark-green header row to an existing table."""
    row = table.add_row()
    cell = row.cells[0]
    if col_span and len(row.cells) > 1:
        for i in range(1, len(row.cells)):
            cell.merge(row.cells[i])
    set_cell_bg(cell, C_DARK_GREEN)
    set_cell_margins(cell, top=50, bottom=50, left=100, right=80)
    add_paragraph_in_cell(cell, text, bold=True, size_pt=9, color=C_WHITE)
    return row


def add_full_width_header(doc, title_text, subtitle_text):
    """Add PROFORMA INVOICE + subtitle banner as a table."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"

    # Title row
    r0 = tbl.rows[0]
    c0 = r0.cells[0]
    set_cell_bg(c0, C_DARK_GREEN)
    set_cell_margins(c0, top=120, bottom=120, left=80, right=80)
    p_title = c0.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)
    add_run(p_title, title_text, bold=True, size_pt=20, color=C_WHITE)

    # Subtitle row
    r1 = tbl.rows[1]
    c1 = r1.cells[0]
    set_cell_bg(c1, C_MID_GREEN)
    set_cell_margins(c1, top=60, bottom=80, left=80, right=80)
    p_sub = c1.paragraphs[0]
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(0)
    add_run(p_sub, subtitle_text, bold=False, size_pt=8, color="C8E6C9")

    return tbl


def add_two_col_info_table(doc, left_rows, right_rows, left_label="SELLER",
                            right_label="BUYER / CONSIGNEE"):
    """
    Creates a 2-column info table with header bars and label/value pairs.
    left_rows / right_rows: list of (label, value) tuples
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"

    # Header cells
    hdr_l = tbl.rows[0].cells[0]
    hdr_r = tbl.rows[0].cells[1]
    set_cell_bg(hdr_l, C_MID_GREEN)
    set_cell_bg(hdr_r, C_MID_GREEN)
    set_cell_margins(hdr_l, top=50, bottom=50, left=80, right=80)
    set_cell_margins(hdr_r, top=50, bottom=50, left=80, right=80)
    add_paragraph_in_cell(hdr_l, left_label,  bold=True, size_pt=9, color=C_WHITE)
    add_paragraph_in_cell(hdr_r, right_label, bold=True, size_pt=9, color=C_WHITE)

    max_rows = max(len(left_rows), len(right_rows))
    for i in range(max_rows):
        row = tbl.add_row()
        cl = row.cells[0]
        cr = row.cells[1]
        set_cell_margins(cl, top=40, bottom=40, left=80, right=80)
        set_cell_margins(cr, top=40, bottom=40, left=80, right=80)

        if i < len(left_rows):
            lbl, val = left_rows[i]
            add_paragraph_in_cell(cl, lbl, bold=True, size_pt=7, color=C_GRAY_TEXT,
                                  space_before=0, space_after=1)
            cl.add_paragraph()
            p_v = cl.paragraphs[-1]
            p_v.paragraph_format.space_before = Pt(0)
            p_v.paragraph_format.space_after  = Pt(3)
            add_run(p_v, val or "—", size_pt=9, color=C_DARK_TEXT)

        if i < len(right_rows):
            lbl, val = right_rows[i]
            add_paragraph_in_cell(cr, lbl, bold=True, size_pt=7, color=C_GRAY_TEXT,
                                  space_before=0, space_after=1)
            cr.add_paragraph()
            p_v = cr.paragraphs[-1]
            p_v.paragraph_format.space_before = Pt(0)
            p_v.paragraph_format.space_after  = Pt(3)
            add_run(p_v, val or "—", size_pt=9, color=C_DARK_TEXT)

    return tbl


# ── MAIN GENERATOR ────────────────────────────────────────────────────────────
def generate_pi_docx(
    pi_number, issue_date_str, validity_days,
    seller_name, seller_address, seller_country,
    seller_email, seller_phone,
    buyer_name, buyer_address, buyer_country,
    buyer_email, buyer_phone,
    fob_data, packaging_type,
    payment_terms, bank_details,
    notes, lang
):
    """
    Generate a Proforma Invoice in DOCX format using python-docx.
    Returns an io.BytesIO buffer containing the .docx file.
    """
    try:
        validity_dt  = datetime.strptime(issue_date_str, "%d %B %Y") + timedelta(days=int(validity_days))
        validity_str = validity_dt.strftime("%d %B %Y")
    except Exception:
        validity_str = f"+{validity_days} days from issue"

    doc = Document()

    # ── PAGE MARGINS ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── HEADER BANNER ─────────────────────────────────────────────────────────
    add_full_width_header(
        doc,
        "PROFORMA INVOICE",
        "InTradeX-Mate  |  A Strategic Initiative from MAGASTU INDOPRIME GROUP (MIG)  |  Indonesian Spice Export"
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── PI META TABLE ─────────────────────────────────────────────────────────
    meta_tbl = doc.add_table(rows=3, cols=2)
    meta_tbl.style = "Table Grid"
    meta_pairs = [
        ("PI Number",   pi_number),
        ("Issue Date",  issue_date_str),
        ("Valid Until", f"{validity_str}  ({validity_days} days)"),
    ]
    for i, (lbl, val) in enumerate(meta_pairs):
        cl = meta_tbl.rows[i].cells[0]
        cv = meta_tbl.rows[i].cells[1]
        remove_cell_borders(cl)
        remove_cell_borders(cv)
        set_cell_margins(cl, top=40, bottom=40, left=60, right=60)
        set_cell_margins(cv, top=40, bottom=40, left=60, right=60)
        add_paragraph_in_cell(cl, lbl, bold=True, size_pt=8, color=C_MID_GREEN)
        add_paragraph_in_cell(cv, val, size_pt=9, color=C_DARK_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── SELLER / BUYER ────────────────────────────────────────────────────────
    seller_rows = [
        ("Company Name", seller_name),
        ("Address",      seller_address),
        ("Country",      seller_country),
        ("Email",        seller_email),
        ("Phone / WhatsApp", seller_phone),
    ]
    buyer_rows = [
        ("Company Name", buyer_name),
        ("Address",      buyer_address),
        ("Country",      buyer_country),
        ("Email",        buyer_email),
        ("Phone / WhatsApp", buyer_phone),
    ]
    add_two_col_info_table(doc, seller_rows, buyer_rows, "SELLER", "BUYER / CONSIGNEE")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── COMMODITY & SHIPMENT DETAILS ──────────────────────────────────────────
    comm_tbl = doc.add_table(rows=1, cols=1)
    comm_tbl.style = "Table Grid"
    section_header_row(comm_tbl, "  COMMODITY & SHIPMENT DETAILS", col_span=4)

    # Replace with 4-col table
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    det_tbl = doc.add_table(rows=4, cols=4)
    det_tbl.style = "Table Grid"
    details = [
        ("Commodity",   fob_data.get("commodity", ""),
         "HS Code",     fob_data.get("hs_code", "")),
        ("Origin",      fob_data.get("origin", ""),
         "Loading Port",fob_data.get("loading_port", "")),
        ("Packaging",   packaging_type,
         "Total Units", str(fob_data.get("total_units_needed", "")) + " unit(s)"),
        ("Net Weight",  f"{fob_data.get('net_weight_kg', 0):,} Kg",
         "Gross Weight",f"{fob_data.get('gross_weight_kg', 0):,} Kg"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(details):
        row = det_tbl.rows[i]
        for j, (txt, is_lbl) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            c = row.cells[j]
            remove_cell_borders(c)
            set_cell_margins(c, top=40, bottom=40, left=60, right=60)
            if is_lbl:
                add_paragraph_in_cell(c, txt, bold=True, size_pt=8, color=C_MID_GREEN)
            else:
                add_paragraph_in_cell(c, txt, size_pt=9, color=C_DARK_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── PRICING TABLE ─────────────────────────────────────────────────────────
    price_hdr_tbl = doc.add_table(rows=1, cols=1)
    price_hdr_tbl.style = "Table Grid"
    section_header_row(
        price_hdr_tbl,
        f"  PRICING  (Incoterm: FOB {fob_data.get('loading_port', '')}, Indonesia)"
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    p_tbl = doc.add_table(rows=3, cols=5)
    p_tbl.style = "Table Grid"

    headers = ["Description", "HS Code", "Qty (Kg)", "Unit Price (USD/Kg)", "Total (USD)"]
    for j, h in enumerate(headers):
        c = p_tbl.rows[0].cells[j]
        set_cell_bg(c, C_MID_GREEN)
        set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        add_paragraph_in_cell(c, h, bold=True, size_pt=8, color=C_WHITE,
                              align=WD_ALIGN_PARAGRAPH.CENTER)

    row_vals = [
        fob_data.get("commodity", ""),
        fob_data.get("hs_code", ""),
        f"{fob_data.get('net_weight_kg', 0):,.0f}",
        f"${fob_data.get('fob_price_per_kg', 0):.4f}",
        f"${fob_data.get('fob_total_usd', 0):,.2f}",
    ]
    for j, val in enumerate(row_vals):
        c = p_tbl.rows[1].cells[j]
        set_cell_margins(c, top=50, bottom=50, left=80, right=80)
        align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        add_paragraph_in_cell(c, val, size_pt=9, color=C_DARK_TEXT, align=align)

    total_row = p_tbl.rows[2]
    for j in range(3):
        c = total_row.cells[j]
        set_cell_bg(c, C_LIGHT_GREEN)
        set_cell_margins(c, top=50, bottom=50, left=80, right=80)
        add_paragraph_in_cell(c, "", size_pt=9)

    c_lbl = total_row.cells[3]
    set_cell_bg(c_lbl, C_LIGHT_GREEN)
    set_cell_margins(c_lbl, top=50, bottom=50, left=80, right=80)
    add_paragraph_in_cell(c_lbl, "TOTAL FOB VALUE", bold=True, size_pt=9,
                          color=C_DARK_GREEN, align=WD_ALIGN_PARAGRAPH.RIGHT)

    c_val = total_row.cells[4]
    set_cell_bg(c_val, C_LIGHT_GREEN)
    set_cell_margins(c_val, top=50, bottom=50, left=80, right=80)
    add_paragraph_in_cell(
        c_val, f"USD {fob_data.get('fob_total_usd', 0):,.2f}",
        bold=True, size_pt=10, color=C_DARK_GREEN, align=WD_ALIGN_PARAGRAPH.RIGHT
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── PAYMENT TERMS + BANK DETAILS ─────────────────────────────────────────
    pt_tbl = doc.add_table(rows=1, cols=2)
    pt_tbl.style = "Table Grid"

    c_pay = pt_tbl.rows[0].cells[0]
    set_cell_margins(c_pay, top=60, bottom=80, left=80, right=80)
    add_paragraph_in_cell(c_pay, "PAYMENT TERMS", bold=True, size_pt=9,
                          color=C_WHITE)
    set_cell_bg(c_pay, C_MID_GREEN)
    c_pay.add_paragraph()
    p_pay = c_pay.paragraphs[-1]
    p_pay.paragraph_format.space_before = Pt(4)
    add_run(p_pay, payment_terms or "—", size_pt=9, color=C_DARK_TEXT)

    c_bank = pt_tbl.rows[0].cells[1]
    set_cell_margins(c_bank, top=60, bottom=80, left=80, right=80)
    add_paragraph_in_cell(c_bank, "BANK DETAILS", bold=True, size_pt=9,
                          color=C_WHITE)
    set_cell_bg(c_bank, C_MID_GREEN)
    bank_text = bank_details if (bank_details and bank_details.strip()) else (
        "[ Bank Name ]\n[ Account Name ]\n[ Account Number ]\n[ Swift / BIC Code ]\n[ Bank Address ]"
    )
    for line in bank_text.split("\n"):
        c_bank.add_paragraph()
        p_b = c_bank.paragraphs[-1]
        p_b.paragraph_format.space_before = Pt(1)
        p_b.paragraph_format.space_after  = Pt(1)
        add_run(p_b, line.strip(), size_pt=9, color=C_DARK_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── NOTES & CONDITIONS ────────────────────────────────────────────────────
    notes_hdr = doc.add_table(rows=1, cols=1)
    notes_hdr.style = "Table Grid"
    section_header_row(notes_hdr, "  NOTES & CONDITIONS")

    default_notes = (
        f"1. PI valid for {validity_days} calendar days from issue date.\n"
        "2. Prices subject to change after validity period.\n"
        "3. Commercial Invoice issued upon order confirmation.\n"
        "4. Phytosanitary Cert, COO, COA available upon request.\n"
        "5. Force majeure clauses apply per international trade practice."
    )
    note_text = notes if (notes and notes.strip()) else default_notes
    for line in note_text.split("\n"):
        if line.strip():
            p_n = doc.add_paragraph()
            p_n.paragraph_format.space_before = Pt(1)
            p_n.paragraph_format.space_after  = Pt(1)
            add_run(p_n, line.strip(), size_pt=8, color=C_GRAY_TEXT)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(8)
    add_run(
        footer_p,
        f"Generated by InTradeX-Mate  |  {pi_number}  |  Issued: {issue_date_str}  |  "
        "This document is computer-generated and valid without a physical stamp unless otherwise stated.",
        size_pt=7, color="888888"
    )

    # ── SAVE TO BUFFER ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
