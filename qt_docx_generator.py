"""
qt_docx_generator.py — Quotation DOCX generator (Python/python-docx)
Replaces the previous Node.js approach.
"""

import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


C_DARK_GREEN  = "1B4332"
C_MID_GREEN   = "2D6A4F"
C_LIGHT_GREEN = "E8F5E9"
C_WHITE       = "FFFFFF"
C_DARK_TEXT   = "1A1A1A"
C_GRAY_TEXT   = "555555"


def set_cell_bg(cell, hex_color: str):
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def remove_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_pr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ))


def add_run(para, text, bold=False, size_pt=9, color=None, italic=False):
    run = para.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_para_in_cell(cell, text, bold=False, size_pt=9, color=None,
                     align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2):
    if len(cell.paragraphs) == 1 and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    add_run(p, text, bold=bold, size_pt=size_pt, color=color)
    return p


def generate_qt_docx(
    qt_number, issue_date_str, valid_until_str, qt_validity,
    qt_seller_name, qt_seller_address, qt_seller_email, qt_seller_phone,
    qt_buyer_name, qt_buyer_company, qt_buyer_email, qt_buyer_country,
    fob_data, qt_payment, qt_notes
):
    """
    Generate a Quotation in DOCX format using python-docx.
    Returns an io.BytesIO buffer containing the .docx file.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(1.8)
        section.right_margin  = Cm(1.8)

    # ── TITLE BANNER ──────────────────────────────────────────────────────────
    banner = doc.add_table(rows=2, cols=1)
    banner.style = "Table Grid"

    r0 = banner.rows[0].cells[0]
    set_cell_bg(r0, C_DARK_GREEN)
    set_cell_margins(r0, top=120, bottom=120)
    add_para_in_cell(r0, "QUOTATION", bold=True, size_pt=22, color=C_WHITE,
                     align=WD_ALIGN_PARAGRAPH.CENTER)

    r1 = banner.rows[1].cells[0]
    set_cell_bg(r1, C_MID_GREEN)
    set_cell_margins(r1, top=60, bottom=80)
    add_para_in_cell(
        r1,
        "InTradeX-Mate  |  A Strategic Initiative from MAGASTU INDOPRIME GROUP (MIG)  |  Indonesian Spice Export",
        size_pt=8, color="C8E6C9", align=WD_ALIGN_PARAGRAPH.CENTER
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── QT META ───────────────────────────────────────────────────────────────
    meta = doc.add_table(rows=2, cols=4)
    meta.style = "Table Grid"
    meta_data = [
        [("QT Number", qt_number), ("Issue Date", issue_date_str)],
        [("Valid Until", f"{valid_until_str} ({qt_validity} days)"), ("", "")],
    ]
    for i, row_data in enumerate(meta_data):
        for j, (lbl, val) in enumerate(row_data):
            c_lbl = meta.rows[i].cells[j * 2]
            c_val = meta.rows[i].cells[j * 2 + 1]
            remove_cell_borders(c_lbl)
            remove_cell_borders(c_val)
            set_cell_margins(c_lbl, top=40, bottom=40, left=60, right=40)
            set_cell_margins(c_val, top=40, bottom=40, left=40, right=60)
            add_para_in_cell(c_lbl, lbl, bold=True, size_pt=9, color=C_MID_GREEN)
            add_para_in_cell(c_val, val, size_pt=9, color=C_DARK_TEXT)

    # Horizontal divider
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after  = Pt(4)
    div_run = div.add_run("─" * 90)
    div_run.font.color.rgb = RGBColor.from_string(C_MID_GREEN)
    div_run.font.size = Pt(6)

    # ── SELLER / BUYER ────────────────────────────────────────────────────────
    sb = doc.add_table(rows=2, cols=2)
    sb.style = "Table Grid"

    hdr_s = sb.rows[0].cells[0]
    hdr_b = sb.rows[0].cells[1]
    set_cell_bg(hdr_s, C_DARK_GREEN)
    set_cell_bg(hdr_b, C_DARK_GREEN)
    set_cell_margins(hdr_s, top=50, bottom=50, left=80, right=80)
    set_cell_margins(hdr_b, top=50, bottom=50, left=80, right=80)
    add_para_in_cell(hdr_s, "SELLER", bold=True, size_pt=9, color=C_WHITE)
    add_para_in_cell(hdr_b, "BUYER / RECIPIENT", bold=True, size_pt=9, color=C_WHITE)

    data_s = sb.rows[1].cells[0]
    data_b = sb.rows[1].cells[1]
    set_cell_margins(data_s, top=60, bottom=60, left=80, right=80)
    set_cell_margins(data_b, top=60, bottom=60, left=80, right=80)

    seller_lines = [
        (qt_seller_name, True, 10),
        (qt_seller_address, False, 9),
        ("Indonesia", False, 9),
        (qt_seller_email, False, 9),
        (qt_seller_phone, False, 9),
    ]
    for (text, bold, sz) in seller_lines:
        if text:
            p = data_s.add_paragraph() if data_s.paragraphs[0].text != "" else data_s.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            col = C_DARK_TEXT if bold else C_GRAY_TEXT
            add_run(p, text, bold=bold, size_pt=sz, color=col)

    buyer_lines = [
        (qt_buyer_name, True, 10),
        (qt_buyer_company, False, 9),
        (qt_buyer_country, False, 9),
        (qt_buyer_email, False, 9),
    ]
    for (text, bold, sz) in buyer_lines:
        if text:
            p = data_b.add_paragraph() if data_b.paragraphs[0].text != "" else data_b.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            col = C_DARK_TEXT if bold else C_GRAY_TEXT
            add_run(p, text, bold=bold, size_pt=sz, color=col)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── COMMODITY & PRICING HEADER ────────────────────────────────────────────
    price_hdr = doc.add_table(rows=1, cols=1)
    price_hdr.style = "Table Grid"
    r = price_hdr.rows[0].cells[0]
    set_cell_bg(r, C_DARK_GREEN)
    set_cell_margins(r, top=50, bottom=50, left=80, right=80)
    add_para_in_cell(r, "  COMMODITY & PRICING DETAILS", bold=True, size_pt=9, color=C_WHITE)

    # ── PRICING TABLE ─────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    pt = doc.add_table(rows=3, cols=5)
    pt.style = "Table Grid"
    hdr_labels = ["Description", "HS Code", "Qty (Kg)", "Unit Price\n(USD/Kg)", "Total (USD)"]
    for j, h in enumerate(hdr_labels):
        c = pt.rows[0].cells[j]
        set_cell_bg(c, C_LIGHT_GREEN)
        set_cell_margins(c, top=60, bottom=60, left=80, right=80)
        add_para_in_cell(c, h, bold=True, size_pt=9,
                         align=WD_ALIGN_PARAGRAPH.CENTER)

    vals = [
        fob_data.get("commodity", ""),
        fob_data.get("hs_code", ""),
        f"{fob_data.get('volume_kg', 0):,.0f}",
        f"${fob_data.get('fob_price_per_kg', 0):.4f}",
        f"${fob_data.get('fob_total_usd', 0):,.2f}",
    ]
    for j, v in enumerate(vals):
        c = pt.rows[1].cells[j]
        set_cell_margins(c, top=50, bottom=50, left=80, right=80)
        align = WD_ALIGN_PARAGRAPH.RIGHT if j >= 2 else WD_ALIGN_PARAGRAPH.LEFT
        add_para_in_cell(c, v, size_pt=9, color=C_DARK_TEXT, align=align)

    # Total row
    for j in range(3):
        c = pt.rows[2].cells[j]
        set_cell_margins(c, top=50, bottom=50, left=80, right=80)
        add_para_in_cell(c, "", size_pt=9)
    c3 = pt.rows[2].cells[3]
    set_cell_margins(c3, top=50, bottom=50, left=80, right=80)
    add_para_in_cell(c3, "TOTAL FOB VALUE", bold=True, size_pt=9,
                     color=C_DARK_GREEN, align=WD_ALIGN_PARAGRAPH.RIGHT)
    c4 = pt.rows[2].cells[4]
    set_cell_margins(c4, top=50, bottom=50, left=80, right=80)
    add_para_in_cell(c4, f"${fob_data.get('fob_total_usd', 0):,.2f}",
                     bold=True, size_pt=10, color=C_DARK_GREEN,
                     align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── SHIPMENT DETAILS ──────────────────────────────────────────────────────
    ship = doc.add_table(rows=3, cols=4)
    ship.style = "Table Grid"
    ship_data = [
        ("Incoterm",      f"FOB {fob_data.get('loading_port', '')}",
         "Origin",        fob_data.get("origin", "")),
        ("Net Weight",    f"{fob_data.get('volume_kg', 0):,.0f} Kg",
         "Gross Weight",  f"{fob_data.get('gross_weight_kg', 0):,.0f} Kg"),
        ("Payment Terms", qt_payment,
         "Price Valid",   f"{valid_until_str} ({qt_validity} days)"),
    ]
    for i, (l1, v1, l2, v2) in enumerate(ship_data):
        for j, (txt, is_lbl) in enumerate([(l1, True), (v1, False), (l2, True), (v2, False)]):
            c = ship.rows[i].cells[j]
            remove_cell_borders(c)
            set_cell_margins(c, top=40, bottom=40, left=60, right=60)
            if is_lbl:
                add_para_in_cell(c, txt, bold=True, size_pt=9, color=C_MID_GREEN)
            else:
                add_para_in_cell(c, txt, size_pt=9, color=C_DARK_TEXT)

    # ── NOTES ─────────────────────────────────────────────────────────────────
    if qt_notes and qt_notes.strip():
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
        nh = doc.add_table(rows=1, cols=1)
        nh.style = "Table Grid"
        r = nh.rows[0].cells[0]
        set_cell_bg(r, C_DARK_GREEN)
        set_cell_margins(r, top=50, bottom=50, left=80, right=80)
        add_para_in_cell(r, "  NOTES & CONDITIONS", bold=True, size_pt=9, color=C_WHITE)

        for line in qt_notes.split("\n"):
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(1)
                add_run(p, line.strip(), size_pt=9, color=C_GRAY_TEXT)

    # ── FOOTER ────────────────────────────────────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        fp,
        f"Generated by InTradeX-Mate  |  {qt_number}  |  Issued: {issue_date_str}",
        size_pt=7, color="888888"
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
